from typing import Tuple, Optional
from io import BytesIO
import os

from docx import Document
from openai import AsyncOpenAI

from utils.file_converter import convert_to_pdf

# Content types
DOCX_MT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MT  = "application/pdf"

# ---------------- Skill Grouping ----------------
def _group_skills(skills: list[str]) -> list[dict]:
    """
    Auto-group a flat list of skills into common categories.
    """
    groups = {
        "Frontend": ["html", "css", "javascript", "typescript", "react", "angular", "vue"],
        "Backend": ["java", "python", "go", "c#", "node", "php", "spring", "django", "flask"],
        "Frameworks": ["spring boot", "express", "hibernate", "jpa", "next.js", "fastapi"],
        "Databases": ["mysql", "postgres", "mongodb", "oracle", "sqlite", "redis", "dynamodb"],
        "Tools": ["git", "docker", "kubernetes", "maven", "gradle", "postman", "vscode", "intellij"],
        "Concepts": ["oop", "object-oriented", "microservices", "algorithms", "data structures", "ci/cd", "agile"],
    }

    result = []
    if not skills:
        return result

    for label, keywords in groups.items():
        found = [s for s in skills if any(k in s.lower() for k in keywords)]
        if found:
            result.append({"label": label, "items": found})

    # Leftover uncategorized skills
    categorized = {s for g in result for s in g["items"]}
    leftover = [s for s in skills if s not in categorized]
    if leftover:
        result.append({"label": "Other", "items": leftover})

    return result


# ---------------- Resume Builder ----------------
async def build_resume(
    resume: dict,
    jd: dict,
    format: str = "docx",
    client: Optional[AsyncOpenAI] = None
) -> bytes:
    """
    Build a tailored resume as DOCX in memory.
    If format='pdf', convert the DOCX to PDF via utils.file_converter.convert_to_pdf.

    Skills are grouped into categories (Frontend, Backend, Frameworks, Databases, Tools, Concepts).
    """

    # Defensive defaults
    if not isinstance(resume, dict):
        resume = {}
    if not isinstance(jd, dict):
        jd = {}

    # --- Build DOCX in memory ---
    doc = Document()

    # Header
    name = resume.get("name", "Candidate Name")
    email = resume.get("email", "")
    phone = resume.get("phone", "")
    linkedin = resume.get("linkedin", "")
    github = resume.get("github", "")
    location = resume.get("location", "")

    doc.add_heading(name, level=0)

    header_line_parts = [p for p in [email, phone, linkedin, github, location] if p]
    if header_line_parts:
        doc.add_paragraph(" | ".join(header_line_parts))

    # Summary
    summary = resume.get("summary") or "Summary available upon request."
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summary)

    # Skills (grouped)
    skills = resume.get("skills", []) or []
    skills_groups = _group_skills(skills)

    if skills_groups:
        doc.add_heading("Skills", level=1)
        for g in skills_groups:
            doc.add_paragraph(f"{g['label']}: {', '.join(g['items'])}", style="List Bullet")

    # Experience
    experience = resume.get("experience", []) or []
    if experience:
        doc.add_heading("Experience", level=1)
        for exp in experience:
            title = exp.get("title", "")
            company = exp.get("company", "")
            dates = exp.get("dates", "")
            header = " - ".join([p for p in [title, company] if p])
            if dates:
                header = f"{header} ({dates})" if header else dates
            if header:
                doc.add_paragraph(header)
            for b in exp.get("bullets", []) or []:
                doc.add_paragraph(str(b), style="List Bullet")

    # Education
    education = resume.get("education", []) or []
    if education:
        doc.add_heading("Education", level=1)
        for edu in education:
            degree = edu.get("degree", "")
            school = edu.get("school", "")
            dates = edu.get("dates", "")
            line = " - ".join([p for p in [degree, school] if p])
            if dates:
                line = f"{line} ({dates})" if line else dates
            if line:
                doc.add_paragraph(line)
            details = edu.get("details", [])
            # optional details list support
            all_details =  details
            doc.add_paragraph(all_details, style="List Bullet")


    # Projects 
    projects = resume.get("projects", []) or []
    if projects:
        doc.add_heading("Projects", level=1)
        for p in projects:
            name_ = p.get("name", "")
            dates = p.get("dates", "")
            header = name_ if name_ else ""
            if dates:
                header = f"{header} ({dates})" if header else dates
            if header:
                doc.add_paragraph(header)
            desc = p.get("description", "")
            if desc:
                doc.add_paragraph(desc)
            for b in p.get("bullets", []) or []:
                doc.add_paragraph(str(b), style="List Bullet")

    # Achievements (optional)
    achievements = resume.get("achievements", []) or []
    if achievements:
        doc.add_heading("Achievements", level=1)
        for a in achievements:
            doc.add_paragraph(str(a), style="List Bullet")

    # Save DOCX to memory
    bio = BytesIO()
    doc.save(bio)
    docx_bytes = bio.getvalue()

    # Convert to PDF if requested
    if str(format).lower() == "pdf":
        return await convert_to_pdf(docx_bytes)

    return docx_bytes


# ---------------- Shim for backward compatibility ----------------
async def update_resume_to_match_jd(
    parsed_resume: dict,
    parsed_jd: dict,
    original_path: str,
    original_ext: str,
    original_content_type: Optional[str],
    client: Optional[AsyncOpenAI]
) -> Tuple[bytes, str, str]:
    """
    Compatibility shim.
    Produces DOCX or PDF with grouped skills.
    """
    ext = (original_ext or "").lower()
    if ext == ".pdf":
        fmt = "pdf"; media_type = PDF_MT; out_ext = ".pdf"
    else:
        fmt = "docx"; media_type = DOCX_MT; out_ext = ".docx"

    file_bytes = await build_resume(parsed_resume, parsed_jd, format=fmt, client=client)
    return file_bytes, media_type, out_ext
